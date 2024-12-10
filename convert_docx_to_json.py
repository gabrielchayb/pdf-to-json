import json
from docx import Document


def docx_to_json(docx_path, output_json_path):
    # Abre o arquivo .docx
    document = Document(docx_path)
    doc_data = {"paragraphs": [], "tables": []}

    # Extrai texto dos parágrafos
    for paragraph in document.paragraphs:
        if paragraph.text.strip():  # Ignorar parágrafos vazios
            doc_data["paragraphs"].append(paragraph.text.strip())

    # Extrai texto das tabelas
    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        doc_data["tables"].append(table_data)

    # Salva em um arquivo JSON
    with open(output_json_path, 'w', encoding='utf-8') as json_file:
        json.dump(doc_data, json_file, ensure_ascii=False, indent=4)

    print(f"Arquivo DOCX convertido para JSON e salvo em {output_json_path}")

# Exemplo de uso
docx_to_json("sdr-info.docx", "output.json")
